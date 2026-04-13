#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def agregar_linea_horizontal(doc):
    """Agrega una línea horizontal"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def convertir_markdown_a_word(archivo_md, archivo_docx):
    """Convierte un archivo Markdown a Word"""
    
    # Leer el archivo markdown
    with open(archivo_md, 'r', encoding='utf-8') as f:
        contenido = f.read()
    
    # Crear documento Word
    doc = Document()
    
    # Establecer márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Dividir el contenido en líneas
    lineas = contenido.split('\n')
    
    i = 0
    while i < len(lineas):
        linea = lineas[i]
        
        # Títulos principales (# )
        if linea.startswith('# '):
            titulo = linea.replace('# ', '').strip()
            p = doc.add_heading(titulo, level=1)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(128, 0, 128)  # Purpura
                run.font.bold = True
        
        # Subtítulos secundarios (## )
        elif linea.startswith('## '):
            subtitulo = linea.replace('## ', '').strip()
            p = doc.add_heading(subtitulo, level=2)
            for run in p.runs:
                run.font.color.rgb = RGBColor(128, 0, 128)
                run.font.bold = True
        
        # Subtítulos terciarios (### )
        elif linea.startswith('### '):
            subtitulo = linea.replace('### ', '').strip()
            p = doc.add_heading(subtitulo, level=3)
            for run in p.runs:
                run.font.bold = True
        
        # Línea separadora (---)
        elif linea.strip() == '---':
            agregar_linea_horizontal(doc)
        
        # Listas con (- )
        elif linea.strip().startswith('- '):
            item = linea.replace('- ', '').strip()
            p = doc.add_paragraph(item, style='List Bullet')
        
        # Listas numeradas (1. 2. 3.)
        elif re.match(r'^\d+\.\s', linea.strip()):
            item = re.sub(r'^\d+\.\s', '', linea.strip())
            p = doc.add_paragraph(item, style='List Number')
        
        # Tablas (detectar markdown table)
        elif linea.strip().startswith('|'):
            # Recopilar toda la tabla
            tabla_lineas = []
            j = i
            while j < len(lineas) and lineas[j].strip().startswith('|'):
                tabla_lineas.append(lineas[j])
                j += 1
            
            if tabla_lineas:
                # Procesar tabla
                headers = [h.strip() for h in tabla_lineas[0].split('|')[1:-1]]
                
                # Crear tabla en Word
                tabla = doc.add_table(rows=len(tabla_lineas) - 1, cols=len(headers))
                tabla.style = 'Light Grid Accent 1'
                
                # Encabezados
                header_cells = tabla.rows[0].cells
                for idx, header in enumerate(headers):
                    header_cells[idx].text = header
                    # Dar formato a encabezados
                    for paragraph in header_cells[idx].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(128, 0, 128)
                
                # Filas de datos
                for row_idx in range(2, len(tabla_lineas)):
                    row_data = [d.strip() for d in tabla_lineas[row_idx].split('|')[1:-1]]
                    row_cells = tabla.rows[row_idx - 1].cells
                    for col_idx, data in enumerate(row_data):
                        row_cells[col_idx].text = data
                
                i = j - 1
        
        # Bloques de código (```python)
        elif linea.strip().startswith('```'):
            # Recopilar bloque de código
            j = i + 1
            lineas_codigo = []
            while j < len(lineas) and not lineas[j].strip().startswith('```'):
                lineas_codigo.append(lineas[j])
                j += 1
            
            if lineas_codigo:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                for linea_cod in lineas_codigo:
                    run = p.add_run(linea_cod + '\n')
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(64, 64, 64)
                
                i = j
        
        # Párrafos normales con negritas y cursivas
        elif linea.strip() and not linea.strip().startswith('|'):
            p = doc.add_paragraph()
            
            # Procesar formato dentro del párrafo
            texto = linea.strip()
            
            # Reemplazar **texto** con negritas
            patron_bold = r'\*\*(.+?)\*\*'
            patron_italic = r'\*(.+?)\*'
            patron_codigo = r'`(.+?)`'
            
            # Simplificar: solo agregar el texto
            if texto:
                # Procesar negritas
                partes = re.split(r'(\*\*.+?\*\*)', texto)
                for parte in partes:
                    if parte.startswith('**') and parte.endswith('**'):
                        run = p.add_run(parte.replace('**', ''))
                        run.bold = True
                    else:
                        p.add_run(parte)
        
        i += 1
    
    # Guardar documento
    doc.save(archivo_docx)
    print(f"✅ Documento creado: {archivo_docx}")

# Ejecutar conversión
if __name__ == '__main__':
    archivo_md = r'c:\LYP\purpura\MANUAL_TECNICO.md'
    archivo_docx = r'c:\LYP\purpura\MANUAL_TECNICO.docx'
    
    convertir_markdown_a_word(archivo_md, archivo_docx)
    print(f"📄 Archivo Word generado exitosamente")
